"""Dev-only utility: add a sample PDF and DOCX to CASE-0001 so the extraction
pipeline (extract_text.py) can be exercised against every supported format
(PDF / DOCX / TXT / CSV). Not part of the ingestion pipeline itself.

Requires the dev extra:
    python -m pip install -r requirements-dev.txt

Usage:
    python src/_generate_sample_docs.py
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from config import settings  # noqa: E402


def make_pdf(path: Path) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    lines = [
        "Motor Vehicle Record (MVR) Report",
        "",
        "Applicant: Jane Doe",
        "License status: Valid, Class D",
        "Violations (past 3 years): None",
        "Accidents (past 3 years): None",
        "Report date: 2025-05-10",
    ]
    for line in lines:
        pdf.cell(0, 8, text=line, new_x="LMARGIN", new_y="NEXT")
    pdf.output(str(path))


def make_docx(path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_heading("Prior Underwriting History", level=1)
    document.add_paragraph("Applicant: Jane Doe")
    document.add_paragraph(
        "Previous policy application (2021): Approved, standard rate. "
        "No adverse findings at that time."
    )
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Policy Number"
    table.rows[0].cells[1].text = "POL-2021-4821"
    table.rows[1].cells[0].text = "Decision"
    table.rows[1].cells[1].text = "Approved (standard)"
    document.save(str(path))


def main() -> int:
    case_dir = settings.RAW_DOCS_DIR / "CASE-0001"
    case_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = case_dir / "mvr_report.pdf"
    docx_path = case_dir / "prior_underwriting_history.docx"

    make_pdf(pdf_path)
    make_docx(docx_path)

    print(f"Created: {pdf_path}")
    print(f"Created: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
