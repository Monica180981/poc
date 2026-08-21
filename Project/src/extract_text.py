"""Extract text from case documents (PDF / Word / plain text), tag each chunk
with case ID and document type, and produce one consolidated case bundle per
case.

Pure technical ingestion — no underwriting rules, no risk scoring, no
decisioning. This just gets clean, tagged text ready for later pipeline
stages (summarization, extraction of key fields, etc.).

Supported inputs: .pdf, .docx, .txt, .md, .csv, .json
Not yet supported: scanned images (OCR is a separate, later capability per
the POC description). Unsupported or unreadable files are recorded with a
non-"ok" extraction_status rather than stopping the whole run.

Usage:
    python src/extract_text.py                   # process every case in raw_docs/
    python src/extract_text.py --case CASE-0001   # process a single case
    python src/extract_text.py --max-chars 2000   # override chunk size
"""

from __future__ import annotations

import argparse
import datetime as dt
import json
import logging
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from config import settings  # noqa: E402

logger = logging.getLogger("extract_text")


# --- Per-format text extraction ---------------------------------------------

def extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())


def extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def extract_plain_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".txt": extract_plain_text,
    ".md": extract_plain_text,
    ".csv": extract_plain_text,
    ".json": extract_plain_text,
}


# --- Chunking ----------------------------------------------------------------

def chunk_text(text: str, max_chars: int = settings.CHUNK_MAX_CHARS) -> list[str]:
    """Split text into chunks on paragraph boundaries, up to ~max_chars each.

    A paragraph longer than max_chars is hard-split so no chunk is unbounded.
    """
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        if len(para) > max_chars:
            if current:
                chunks.append(current)
                current = ""
            for i in range(0, len(para), max_chars):
                chunks.append(para[i : i + max_chars])
            continue

        if current and len(current) + len(para) + 2 > max_chars:
            chunks.append(current)
            current = para
        else:
            current = f"{current}\n\n{para}" if current else para

    if current:
        chunks.append(current)
    return chunks


# --- Per-document / per-case processing --------------------------------------

def process_document(path: Path, case_id: str, max_chars: int) -> tuple[dict, list[dict]]:
    """Extract, chunk, and tag one document. Never raises — failures are
    recorded on the returned entry so one bad file doesn't stop the batch."""
    doc_type = settings.doc_type_hint(path.name)
    extractor = EXTRACTORS.get(path.suffix.lower())

    entry = {
        "name": path.name,
        "document_type": doc_type,
        "extraction_status": "ok",
        "chunk_count": 0,
    }

    if extractor is None:
        entry["extraction_status"] = "unsupported"
        entry["note"] = (
            f"No extractor for '{path.suffix}' "
            "(e.g. scanned images need OCR — not yet implemented)"
        )
        return entry, []

    try:
        text = extractor(path)
    except Exception as exc:  # noqa: BLE001 - isolate failures per document
        logger.warning("Failed to extract %s: %s", path.name, exc)
        entry["extraction_status"] = "error"
        entry["note"] = str(exc)
        return entry, []

    if not text.strip():
        entry["extraction_status"] = "empty"
        entry["note"] = "No extractable text (possibly a scanned/image-only document)"
        return entry, []

    pieces = chunk_text(text, max_chars=max_chars)
    chunks = [
        {
            "chunk_id": f"{case_id}__{path.name}__{i:03d}",
            "case_id": case_id,
            "document": path.name,
            "document_type": doc_type,
            "chunk_index": i,
            "char_count": len(piece),
            "text": piece,
        }
        for i, piece in enumerate(pieces)
    ]
    entry["chunk_count"] = len(chunks)
    return entry, chunks


def process_case(case_dir: Path, max_chars: int) -> dict:
    """Build the consolidated bundle for one case: every document's status
    plus every tagged chunk, in one dict ready to serialize as JSON."""
    case_id = case_dir.name
    documents: list[dict] = []
    all_chunks: list[dict] = []

    for path in sorted(case_dir.rglob("*")):
        if not path.is_file() or settings.is_ignorable_file(path.name):
            continue
        entry, chunks = process_document(path, case_id, max_chars)
        documents.append(entry)
        all_chunks.extend(chunks)

    return {
        "case_id": case_id,
        "generated_at": dt.datetime.now().isoformat(timespec="seconds"),
        "documents": documents,
        "chunk_count": len(all_chunks),
        "chunks": all_chunks,
    }


def log_bundle(bundle: dict, out_path: Path) -> None:
    logger.info(
        "%s: %d document(s), %d chunk(s) -> %s",
        bundle["case_id"], len(bundle["documents"]), bundle["chunk_count"], out_path,
    )
    for doc in bundle["documents"]:
        flag = "" if doc["extraction_status"] == "ok" else f"  [{doc['extraction_status']}]"
        logger.info(
            "  %-32s %-40s %d chunk(s)%s",
            doc["name"], doc["document_type"], doc["chunk_count"], flag,
        )


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--root", type=Path, default=settings.RAW_DOCS_DIR,
        help="Directory containing per-case sub-folders (default: raw_docs/).",
    )
    parser.add_argument(
        "--out", type=Path, default=settings.PROCESSED_TEXT_DIR,
        help="Directory to write per-case bundle JSON files (default: processed_text/).",
    )
    parser.add_argument("--case", help="Process only this case ID (folder name).")
    parser.add_argument(
        "--max-chars", type=int, default=settings.CHUNK_MAX_CHARS,
        help="Max characters per chunk (default: %(default)s).",
    )
    args = parser.parse_args(argv)

    logging.basicConfig(
        level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s", datefmt="%H:%M:%S"
    )

    if not args.root.exists():
        logger.error("raw_docs directory not found: %s", args.root)
        return 1

    case_dirs = sorted(p for p in args.root.iterdir() if p.is_dir())
    if args.case:
        case_dirs = [d for d in case_dirs if d.name == args.case]
        if not case_dirs:
            logger.error("Case '%s' not found under %s", args.case, args.root)
            return 1

    if not case_dirs:
        logger.warning("No case folders found under %s", args.root)
        return 0

    args.out.mkdir(parents=True, exist_ok=True)

    for case_dir in case_dirs:
        bundle = process_case(case_dir, args.max_chars)
        out_path = args.out / f"{bundle['case_id']}.json"
        out_path.write_text(json.dumps(bundle, indent=2), encoding="utf-8")
        log_bundle(bundle, out_path)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
